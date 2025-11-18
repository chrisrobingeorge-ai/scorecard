# docx_utils.py

from __future__ import annotations

from io import BytesIO
from typing import Dict, Any
from datetime import datetime
import json
import re

import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app_config import OBJECTIVES_DF


JSON_PREFIX = "AB_SCORECARD_JSON:"

# ─────────────────────────────────────────────────────────────────────────────
# Helper text utilities
# ─────────────────────────────────────────────────────────────────────────────
def _set_margins(
    doc: Document,
    top: float = 0.75,
    bottom: float = 0.75,
    left: float = 0.75,
    right: float = 0.75,
):
    """
    Set page size and margins (in inches) for all sections in the document.
    Forces Letter size so all scorecard docs are identical.
    """
    from docx.shared import Inches

    for section in doc.sections:
        # Force Letter page size (8.5" x 11") for consistency
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def _to_plain_text(val: Any) -> str:
    """
    Normalise various AI output types to a human-readable string.
    """
    if val is None:
        return ""
    if isinstance(val, str):
        return val
    if isinstance(val, list):
        return "\n\n".join(_to_plain_text(v) for v in val)
    if isinstance(val, dict):
        if "text" in val and isinstance(val["text"], str):
            return val["text"]
        try:
            return json.dumps(val, ensure_ascii=False, indent=2)
        except Exception:
            return str(val)
    return str(val)


def _strip_objective_codes(text: str) -> str:
    """
    Remove internal objective codes like ART1 / ART2 / ART3 and
    clean up some odd dash list framing.
    """
    if not text:
        return ""
    # Remove "(ART1)" style codes
    s = re.sub(r"\(ART[0-9]+\)", "", text)
    # Remove bare ART1 / ART2 tokens
    s = re.sub(r"\bART[0-9]+\b", "", s)
    # Replace en-dash bullet fragments " – " with a space
    s = s.replace(" – ", " ")
    # Normalise excess whitespace
    s = re.sub(r"\s{2,}", " ", s)
    return s.strip()


def _split_paragraphs(raw_value: Any) -> list[str]:
    """
    Split narrative into logical paragraphs.

    Priority:
    1) If value is a list: each item -> paragraph.
    2) If string: split on blank lines.
    3) If still one block: split into chunks of ~3 sentences.
    """
    # List case: treat each as paragraph
    if isinstance(raw_value, list):
        paras: list[str] = []
        for item in raw_value:
            txt = _strip_objective_codes(_to_plain_text(item)).strip()
            if txt:
                paras.append(txt)
        return paras

    # String case
    text = _strip_objective_codes(_to_plain_text(raw_value or "")).replace("\r\n", "\n")
    if not text.strip():
        return []

    # Try blank-line split first
    parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(parts) > 1:
        return parts

    # If we have single newlines, allow those to break paragraphs
    if "\n" in text:
        parts = [p.strip() for p in text.split("\n") if p.strip()]
        if len(parts) > 1:
            return parts

    # Fallback: sentence-based chunking (group ~3 sentences per paragraph)
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    if len(sentences) <= 3:
        return [" ".join(sentences)] if sentences else []

    paras: list[str] = []
    chunk: list[str] = []
    for s in sentences:
        chunk.append(s)
        if len(chunk) >= 3:
            paras.append(" ".join(chunk))
            chunk = []
    if chunk:
        paras.append(" ".join(chunk))
    return paras


# ─────────────────────────────────────────────────────────────────────────────
# Score utilities
# ─────────────────────────────────────────────────────────────────────────────
def _parse_score_hint(score_hint: str | None) -> float | None:
    """Extract an approximate 0–3 score from a score_hint string."""
    if not score_hint:
        return None

    text = str(score_hint)
    frac_match = re.search(r"(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?)", text)
    if frac_match:
        num = float(frac_match.group(1))
        den = float(frac_match.group(2))
        if den != 0:
            return max(0.0, min(3.0, (num / den) * 3))

    num_match = re.search(r"(\d+(?:\.\d+)?)", text)
    if num_match:
        value = float(num_match.group(1))
        return max(0.0, min(3.0, value))

    return None


def _score_to_color(score: float | None) -> RGBColor:
    """Return a background colour based on score performance."""
    if score is None:
        return RGBColor(204, 209, 217)  # muted grey-blue
    if score >= 2.5:
        return RGBColor(102, 171, 161)  # teal/green
    if score >= 1.5:
        return RGBColor(237, 186, 99)  # amber
    return RGBColor(232, 112, 82)  # coral red


def _score_display(score: float | None) -> str:
    if score is None:
        return "N/A"
    return f"{score:.2f}"


def _set_cell_background(cell, color: RGBColor):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), f"{color.rgb:06X}")
    cell._tc.get_or_add_tcPr().append(shading_elm)


# ─────────────────────────────────────────────────────────────────────────────
# Tables
# ─────────────────────────────────────────────────────────────────────────────
def _responses_table(doc: Document, questions: pd.DataFrame, responses: Dict[str, Any]):
    """
    Build a 3-column table of context / question / response with wrapped text.
    Uses landscape orientation for better fit.
    """
    # Change section to landscape orientation for the responses table
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Pillar / Production / Metric"
    header_cells[1].text = "Question"
    header_cells[2].text = "Response"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for _, row in questions.sort_values("display_order").iterrows():
        qid = row["question_id"]
        qid_str = str(qid)

        prod_label = row.get("production_title", "") or row.get("production", "")
        label = f"{row.get('strategic_pillar', '')} / {prod_label} / {row.get('metric', '')}"
        
        # Get the question text
        question_text = row.get("question_text", "") or row.get("metric", "") or ""
        
        raw_val = responses.get(qid_str, responses.get(qid, ""))

        if isinstance(raw_val, dict):
            primary = raw_val.get("primary", "")
            desc = raw_val.get("description", "")
            if desc:
                value_str = f"{primary} — {desc}"
            else:
                value_str = str(primary)
        else:
            value_str = str(raw_val)

        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = question_text
        row_cells[2].text = value_str
        
        # Format cells
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def build_strategic_index_appendix(doc: Document):
    """
    Build an appendix section showing the strategic objectives index,
    excluding the 'plan_anchor' column.
    """
    doc.add_page_break()
    
    heading = doc.add_heading('Appendix A — Strategic Objectives Index', level=2)
    
    df = OBJECTIVES_DF.copy()

    # If nothing is configured, return a simple notice instead
    if df.empty:
        doc.add_paragraph("No strategic objectives index is currently configured.")
        return

    # Drop plan_anchor if present
    if "plan_anchor" in df.columns:
        df = df.drop(columns=["plan_anchor"])

    # Reorder columns for readability, only if they exist
    preferred_order = ["owner", "objective_id", "objective_title", "short_description"]
    columns = [c for c in preferred_order if c in df.columns]
    if columns:
        df = df[columns]

    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = col.replace("_", " ").title()
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row.tolist()):
            text = "" if val is None else str(val)
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


# ─────────────────────────────────────────────────────────────────────────────
# Main DOCX builder
# ─────────────────────────────────────────────────────────────────────────────
def build_scorecard_docx(
    meta: Dict[str, Any],
    questions: pd.DataFrame,
    responses: Dict[str, Any],
    ai_result: Dict[str, Any],
    logo_path: str | None = None,
) -> bytes:
    """
    Build and return the DOCX as raw bytes.

    - Preserves your "Strategic Summary Scorecard" layout (header, executive summary,
      pillars, by production, risks/priorities, raw responses).
    """
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ─────────────────────────────────────────────────────────────────────
    # AI summary text
    # ─────────────────────────────────────────────────────────────────────
    overall_value = ai_result.get("overall_summary", "") or ""
    ai_summary_text = _to_plain_text(overall_value)

    # ─────────────────────────────────────────────────────────────────────
    # Derive scores from objective_summaries (or pillar_summaries for backward compatibility)
    # ─────────────────────────────────────────────────────────────────────
    objective_summaries = [ps for ps in (ai_result.get("objective_summaries") or ai_result.get("pillar_summaries") or []) if isinstance(ps, dict)]

    score_values: list[float] = []
    pillar_score_map: Dict[str, float] = {}
    objective_score_map: Dict[str, float] = {}

    for obj_sum in objective_summaries:
        # Handle both old pillar structure and new objective structure
        obj_id = obj_sum.get("objective_id", "")
        pillar_key = _to_plain_text(obj_sum.get("strategic_pillar", "")).strip()
        score_val = _parse_score_hint(obj_sum.get("score_hint", ""))

        if score_val is not None:
            score_values.append(score_val)

        # Map scores by both objective_id and pillar name for compatibility
        if obj_id and score_val is not None:
            objective_score_map[obj_id] = score_val
        if pillar_key and score_val is not None:
            pillar_score_map[pillar_key] = score_val

    total_score = sum(score_values) / len(score_values) if score_values else None

    # ─────────────────────────────────────────────────────────────────────
    # Header: logo, title, total score
    # ─────────────────────────────────────────────────────────────────────
    reporting_period = meta.get("month", "")
    department = meta.get("department", "") or "—"

    # Add logo if available
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(2.0))
        except Exception:
            pass
    
    # Title
    title = doc.add_heading('Strategic Summary Scorecard', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Department and reporting period
    p = doc.add_paragraph()
    p.add_run('DEPARTMENT: ').bold = True
    p.add_run(str(department))
    
    p = doc.add_paragraph()
    p.add_run('REPORTING PERIOD: ').bold = True
    p.add_run(str(reporting_period) if reporting_period else "—")
    
    # Total score
    total_score_display = _score_display(total_score)
    total_value = f"{total_score_display} / 3" if total_score_display != "N/A" else "Not rated"
    
    p = doc.add_paragraph()
    p.add_run('TOTAL SCORE: ').bold = True
    run = p.add_run(total_value)
    run.bold = True
    run.font.size = Pt(14)
    
    # Add color background to score
    if total_score is not None:
        total_colour = _score_to_color(total_score)
        run.font.color.rgb = RGBColor(255, 255, 255)  # white text
    
    doc.add_paragraph()  # spacing

    # ─────────────────────────────────────────────────────────────────────
    # Executive summary
    # ─────────────────────────────────────────────────────────────────────
    paragraphs = _split_paragraphs(overall_value)

    if paragraphs:
        doc.add_heading('Executive Summary', level=2)
        for para in paragraphs:
            p = doc.add_paragraph(para)
            p.style = 'Normal'
        doc.add_paragraph()  # spacing

    # ─────────────────────────────────────────────────────────────────────
    # Strategic objectives – narrative
    # ─────────────────────────────────────────────────────────────────────
    if objective_summaries:
        # For School, make the section heading explicitly about the three streams
        dept_lower = str(department).lower()
        if "school" in dept_lower:
            section_label = (
                "School Streams "
                "(Classical Training / Attracting Students / Student Accessibility)"
            )
        else:
            section_label = "Strategic Objectives"

        doc.add_heading(section_label, level=2)

        for obj_sum in objective_summaries:
            # Handle both old pillar structure and new objective structure
            obj_id = obj_sum.get("objective_id", "")
            obj_title = obj_sum.get("objective_title", "") or obj_sum.get("strategic_pillar", "") or "Objective"
            
            # Create display name with ID if available
            if obj_id:
                objective_name_raw = f"{obj_id}: {obj_title}"
            else:
                objective_name_raw = obj_title
            objective_name = _strip_objective_codes(_to_plain_text(objective_name_raw).strip())

            score_hint_raw = _to_plain_text(obj_sum.get("score_hint", "")).strip()
            score_hint_clean = _strip_objective_codes(score_hint_raw)

            summary_text_raw = _to_plain_text(obj_sum.get("summary", "")).strip()
            summary_text = _strip_objective_codes(summary_text_raw)

            score_value = _parse_score_hint(score_hint_raw)
            score_str = _score_display(score_value)
            if score_str != "N/A":
                heading_text = f"{objective_name} — {score_str} / 3"
            else:
                heading_text = objective_name

            if score_hint_clean:
                heading_text = f"{heading_text} ({score_hint_clean})"

            doc.add_heading(heading_text, level=3)
            if summary_text:
                for p in _split_paragraphs(summary_text):
                    doc.add_paragraph(p)
            else:
                doc.add_paragraph("No narrative summary provided for this objective.")

        doc.add_paragraph()  # spacing

    # ─────────────────────────────────────────────────────────────────────
    # By Production / Programme
    # ─────────────────────────────────────────────────────────────────────
    production_summaries = ai_result.get("production_summaries", []) or []
    if production_summaries:
        doc.add_heading('By Production / Programme', level=2)

        for prod in production_summaries:
            if not isinstance(prod, dict):
                continue

            pname_raw = _to_plain_text(prod.get("production") or "General").strip() or "General"
            pname = _strip_objective_codes(pname_raw)
            doc.add_heading(pname, level=3)

            # Handle both old "pillars" and new "objectives" structure
            objectives = prod.get("objectives") or prod.get("pillars") or []
            summaries: list[str] = []
            for obj in objectives:
                if not isinstance(obj, dict):
                    continue
                summary_text_raw = _to_plain_text(obj.get("summary", "")).strip()
                summary_text = _strip_objective_codes(summary_text_raw)
                if summary_text:
                    summaries.append(summary_text)

            # Join with paragraph breaks to preserve structure, then split properly
            combined = "\n\n".join(summaries).strip()
            if combined:
                combined_paras = _split_paragraphs(combined)
                for p in combined_paras:
                    doc.add_paragraph(p)
            else:
                doc.add_paragraph("No summary provided.")

    # ─────────────────────────────────────────────────────────────────────
    # Risks, priorities, notes
    # ─────────────────────────────────────────────────────────────────────
    risks = [
        _strip_objective_codes(_to_plain_text(r)).strip()
        for r in (ai_result.get("risks", []) or [])
        if _to_plain_text(r).strip()
    ]
    if risks:
        doc.add_heading('Key Risks / Concerns', level=2)
        for r in risks:
            doc.add_paragraph(f"• {r}", style='List Bullet')

    priorities = [
        _strip_objective_codes(_to_plain_text(p)).strip()
        for p in (ai_result.get("priorities_next_month", []) or [])
        if _to_plain_text(p).strip()
    ]
    if priorities:
        doc.add_heading('Priorities for Next Period', level=2)
        for p in priorities:
            doc.add_paragraph(f"• {p}", style='List Bullet')

    nfl_raw = _to_plain_text(ai_result.get("notes_for_leadership", "")).strip()
    nfl = _strip_objective_codes(nfl_raw)
    if nfl:
        doc.add_heading('Notes for Leadership', level=2)
        for p in _split_paragraphs(nfl):
            doc.add_paragraph(p)

    # ─────────────────────────────────────────────────────────────────────
    # Raw responses on a fresh page
    # ─────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading('Raw Scorecard Responses', level=2)
    _responses_table(doc, questions, responses)

    # Note: Strategic Objectives Index is only included in the overall summary,
    # not in individual department scorecards

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes


def build_overall_board_docx(
    reporting_label: str,
    dept_overview: pd.DataFrame,
    ai_result: Dict[str, Any],
    logo_path: str | None = None,
) -> bytes:
    """
    Build a Board-facing DOCX for the overall monthly scorecard.

    - reporting_label: e.g. "November 2025" or "2025-11"
    - dept_overview: DataFrame with at least columns:
        ["department", "month_label", "overall_score"]
    - ai_result: the dict returned by interpret_overall_scorecards(...)
    """
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Header: logo + title ─────────────
    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(2.0))
        except Exception:
            pass
    
    title = doc.add_heading('Strategic Summary Scorecard', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # ── Reporting period ────────────────
    if reporting_label:
        p = doc.add_paragraph()
        p.add_run('REPORTING PERIOD: ').bold = True
        p.add_run(str(reporting_label))
        doc.add_paragraph()  # spacing

    # ─────────────────────────────────────────────────────────────────────
    # Departments overview table
    # ─────────────────────────────────────────────────────────────────────
    if not dept_overview.empty:
        doc.add_heading('Departments included', level=2)

        display_cols = ["department", "month_label", "overall_score"]
        display_cols = [c for c in display_cols if c in dept_overview.columns]

        table = doc.add_table(rows=1, cols=len(display_cols))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(display_cols):
            header_cells[i].text = col.replace("_", " ").title()
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for _, row in dept_overview[display_cols].iterrows():
            row_cells = table.add_row().cells
            for i, c in enumerate(display_cols):
                row_cells[i].text = str(row.get(c, ""))
        
        doc.add_paragraph()  # spacing

    # ─────────────────────────────────────────────────────────────────────
    # Main Board narrative
    # ─────────────────────────────────────────────────────────────────────
    doc.add_heading('Board Narrative', level=2)

    overall_text = (ai_result.get("overall_summary") or "").strip()

    if not overall_text:
        doc.add_paragraph("No Board report text was generated.")
    else:
        text = overall_text.replace("\r\n", "\n")
        parts = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in parts:
            doc.add_paragraph(para)

    # ─────────────────────────────────────────────────────────────────────
    # Strategic Pillar risks / concerns
    # ─────────────────────────────────────────────────────────────────────
    risks = ai_result.get("risks") or []
    if risks:
        doc.add_heading('Strategic Pillar – Risks / Concerns', level=2)
        for r in risks:
            doc.add_paragraph(f"• {str(r)}", style='List Bullet')

    # Organisation-wide priorities
    priorities = ai_result.get("priorities_next_month") or []
    if priorities:
        doc.add_heading('Organisation-wide Priorities for Next Period', level=2)
        for p in priorities:
            doc.add_paragraph(f"• {str(p)}", style='List Bullet')

    # Notes for leadership
    notes = (ai_result.get("notes_for_leadership") or "").strip()
    if notes:
        doc.add_heading('Notes for Leadership', level=2)
        parts = [p.strip() for p in notes.replace("\r\n", "\n").split("\n\n") if p.strip()]
        for para in parts:
            doc.add_paragraph(para)

    # ─────────────────────────────────────────────────────────────────────
    # Appendix A — Strategic Objectives Index
    # ─────────────────────────────────────────────────────────────────────
    build_strategic_index_appendix(doc)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes
